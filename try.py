import speech_recognition as sr
from docx import Document
import tkinter as tk
from tkinter.scrolledtext import ScrolledText
from tkinter import messagebox
import threading

# Initialize the recognizer
recognizer = sr.Recognizer()
stop_listening = False  # Flag to control the loop
recognized_text_storage = ""  # Store only recognized text for saving

# Function to recognize speech and return text
def recognize_speech():
    mic = sr.Microphone()

    try:
        with mic as source:
            recognizer.adjust_for_ambient_noise(source, duration=1)
            recognizer.energy_threshold = 300
            recognizer.pause_threshold = 0.8

            audio = recognizer.listen(source, phrase_time_limit=8)

        text = recognizer.recognize_google(audio, language='kn-IN')
        return text

    except sr.RequestError:
        return "[Error: Unable to reach the speech recognition service]"
    except sr.UnknownValueError:
        return "[Error: Could not understand audio]"
    except Exception as e:
        return f"[Error: {str(e)}]"

# Function to handle speech-to-text continuously until stopped
def continuous_speech_to_text():
    global stop_listening, recognized_text_storage
    while not stop_listening:
        recognized_text = recognize_speech()
        # Only add valid text (not error messages) to recognized_text_storage
        if recognized_text and not recognized_text.startswith("[Error"):
            # Append recognized text to storage and text box
            root.after(0, lambda: text_box.insert(tk.END, recognized_text + " "))
            root.after(0, lambda: text_box.yview(tk.END))
            recognized_text_storage += recognized_text + " "

# Function to start the continuous speech-to-text in a separate thread
def start_listening():
    global stop_listening
    stop_listening = False
    text_box.insert(tk.END, "Started listening...\n")  # Only for UI feedback
    text_box.yview(tk.END)
    threading.Thread(target=continuous_speech_to_text, daemon=True).start()

# Function to stop listening
def stop_listening_func():
    global stop_listening
    stop_listening = True
    text_box.insert(tk.END, "\nStopped listening.\n")  # Only for UI feedback

# Function to save the text to a Word document
def save_to_word():
    global recognized_text_storage
    if recognized_text_storage.strip():
        file_name = file_name_entry.get() or "output"
        doc = Document()
        doc.add_heading('Kannada Speech to Text', 0)
        doc.add_paragraph(recognized_text_storage.strip())  # Save only valid recognized text
        doc.save(f"{file_name}.docx")
        messagebox.showinfo("Success", f"Document saved as {file_name}.docx")
    else:
        messagebox.showwarning("Warning", "No text to save!")

# Creating the main window for UI
root = tk.Tk()
root.title("Kannada Speech to Text")
root.geometry("500x500")

# Label for instructions
label = tk.Label(root, text="Press 'Start' to begin speaking in Kannada. Press 'Stop' to end:", font=("Arial", 12))
label.pack(pady=10)

# Button to start continuous speech-to-text conversion
start_button = tk.Button(root, text="Start Listening", command=start_listening, font=("Arial", 12), bg="lightblue")
start_button.pack(pady=10)

# Button to stop listening
stop_button = tk.Button(root, text="Stop Listening", command=stop_listening_func, font=("Arial", 12), bg="lightcoral")
stop_button.pack(pady=10)

# Scrollable Text box to display the recognized Kannada text
text_box = ScrolledText(root, wrap=tk.WORD, width=60, height=10, font=("Arial", 14), bg="lightyellow")
text_box.pack(pady=10)

# Entry box for the file name
file_name_label = tk.Label(root, text="Enter file name to save the document (optional):", font=("Arial", 10))
file_name_label.pack(pady=5)
file_name_entry = tk.Entry(root, font=("Arial", 12), width=30)
file_name_entry.pack(pady=5)

# Button to save the text as a Word document
save_button = tk.Button(root, text="Save to Word", command=save_to_word, font=("Arial", 12), bg="lightgreen")
save_button.pack(pady=10)

# Run the application
root.mainloop()
